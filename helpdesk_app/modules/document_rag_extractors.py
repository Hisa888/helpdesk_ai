from __future__ import annotations

import io
import re
import unicodedata
from pathlib import Path
from typing import Any, Iterable

from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document

SUPPORTED_DOC_RAG_EXTENSIONS = ("pdf", "docx", "xlsx", "xlsm", "txt", "md")

_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_BOX_RE = re.compile(r"[\u25A0-\u25FF\u2580-\u259F\uE000-\uF8FF]+")
_SPACE_RE = re.compile(r"[ \t　]+")


def normalize_doc_text(text: Any) -> str:
    """PDF/Word/Excel/TXT から取り出した文字列を検索・FAQ生成向けに整える。"""
    if text is None:
        return ""
    s = str(text)
    if s.lower().strip() in {"nan", "none", "null"}:
        return ""
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = _CONTROL_RE.sub("", s)
    # PDFフォント抽出で混入しやすい黒四角・私用領域文字を除去
    s = _BOX_RE.sub("", s).replace("□", "").replace("■", "")
    s = _SPACE_RE.sub(" ", s)
    # 行末の余分な空白を除去
    s = "\n".join(line.strip() for line in s.split("\n"))
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _decode_text_bytes(file_bytes: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp932", "shift_jis", "euc_jp"):
        try:
            return file_bytes.decode(enc)
        except Exception:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def iter_text_chunks(text: str, *, chunk_size: int = 700, overlap: int = 120) -> Iterable[str]:
    clean = normalize_doc_text(text)
    if not clean:
        return
    start = 0
    text_len = len(clean)
    while start < text_len:
        end = min(text_len, start + chunk_size)
        if end < text_len:
            cut = clean.rfind("\n", start, end)
            if cut <= start + 120:
                cut = clean.rfind("。", start, end)
            if cut <= start + 120:
                cut = clean.rfind(" ", start, end)
            if cut > start + 120:
                end = cut + 1
        piece = clean[start:end].strip()
        if piece:
            yield piece
        if end >= text_len:
            break
        start = max(end - overlap, start + 1)


def _looks_like_noise(text: str) -> bool:
    clean = normalize_doc_text(text)
    if not clean:
        return True
    # 文字化け・記号ばかりのページを除外
    jp_alnum = re.findall(r"[A-Za-z0-9ぁ-んァ-ン一-龥]", clean)
    return len(jp_alnum) < max(8, len(clean) * 0.18)


def extract_pdf_sections(file_bytes: bytes, filename: str) -> list[dict[str, str]]:
    """PDFをページ単位で抽出。

    画像スキャンPDFはOCRなしでは本文抽出できないため、その場合は空になります。
    """
    results: list[dict[str, str]] = []
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception:
        return results

    for idx, page in enumerate(reader.pages, start=1):
        texts: list[str] = []
        # extraction_mode が使える pypdf では layout を優先。古い版では通常抽出へフォールバック。
        for mode in ("layout", None):
            try:
                if mode:
                    t = page.extract_text(extraction_mode=mode) or ""
                else:
                    t = page.extract_text() or ""
            except TypeError:
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""
            except Exception:
                t = ""
            t = normalize_doc_text(t)
            if t and not _looks_like_noise(t):
                texts.append(t)
                break
        text = normalize_doc_text("\n".join(texts))
        if not text:
            continue
        results.append({
            "source_name": filename,
            "source_type": "pdf",
            "location": f"page {idx}",
            "text": text,
        })
    return results


def _iter_docx_table_text(doc: Document) -> list[str]:
    blocks: list[str] = []
    for t_idx, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for r_idx, row in enumerate(table.rows, start=1):
            cells = [normalize_doc_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                rows.append(f"row {r_idx}: " + " / ".join(cells))
        if rows:
            blocks.append(f"[表 {t_idx}]\n" + "\n".join(rows))
    return blocks


def extract_docx_sections(file_bytes: bytes, filename: str) -> list[dict[str, str]]:
    """Word(docx)を見出し・段落・表まで含めて抽出。"""
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        return []

    results: list[dict[str, str]] = []
    current_title = "document"
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines, current_title
        text = normalize_doc_text("\n".join(current_lines))
        if text:
            results.append({
                "source_name": filename,
                "source_type": "docx",
                "location": current_title or "document",
                "text": text,
            })
        current_lines = []

    for p in doc.paragraphs:
        raw = normalize_doc_text(p.text)
        if not raw:
            continue
        style_name = (getattr(getattr(p, "style", None), "name", "") or "").lower()
        # 見出しでセクションを分割。日本語スタイル名も考慮。
        if "heading" in style_name or "見出し" in style_name:
            flush()
            current_title = raw[:80]
            current_lines.append(raw)
        else:
            current_lines.append(raw)
    flush()

    # Word内の表もFAQ生成では重要なので別セクションとして追加
    for idx, table_text in enumerate(_iter_docx_table_text(doc), start=1):
        if table_text:
            results.append({
                "source_name": filename,
                "source_type": "docx",
                "location": f"table {idx}",
                "text": table_text,
            })

    # 段落なし・表だけの文書対策
    if not results:
        table_blocks = _iter_docx_table_text(doc)
        text = normalize_doc_text("\n\n".join(table_blocks))
        if text:
            results.append({
                "source_name": filename,
                "source_type": "docx",
                "location": "document",
                "text": text,
            })
    return results


def extract_xlsx_sections(file_bytes: bytes, filename: str) -> list[dict[str, str]]:
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    results: list[dict[str, str]] = []
    for ws in wb.worksheets:
        rows_text: list[str] = []
        for r_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            values = [normalize_doc_text(v) for v in row if normalize_doc_text(v)]
            if values:
                rows_text.append(f"row {r_idx}: " + " / ".join(values))
        text = normalize_doc_text("\n".join(rows_text))
        if text:
            results.append({
                "source_name": filename,
                "source_type": "xlsx",
                "location": f"sheet {ws.title}",
                "text": text,
            })
    return results


def extract_text_sections(file_bytes: bytes, filename: str, source_type: str) -> list[dict[str, str]]:
    text = normalize_doc_text(_decode_text_bytes(file_bytes))
    if not text:
        return []
    return [{
        "source_name": filename,
        "source_type": source_type,
        "location": "document",
        "text": text,
    }]


def extract_sections_from_uploaded_file(uploaded_file: Any) -> list[dict[str, str]]:
    filename = str(getattr(uploaded_file, "name", "document")).strip() or "document"
    ext = Path(filename).suffix.lower().lstrip(".")
    file_bytes = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()

    if ext == "pdf":
        return extract_pdf_sections(file_bytes, filename)
    if ext == "docx":
        return extract_docx_sections(file_bytes, filename)
    if ext in {"xlsx", "xlsm"}:
        return extract_xlsx_sections(file_bytes, filename)
    if ext in {"txt", "md"}:
        return extract_text_sections(file_bytes, filename, ext)
    return []


def build_chunks_from_sections(sections: list[dict[str, str]], *, chunk_size: int = 700, overlap: int = 120) -> list[dict[str, str]]:
    chunks: list[dict[str, str]] = []
    for sec in sections:
        pieces = list(iter_text_chunks(sec.get("text", ""), chunk_size=chunk_size, overlap=overlap))
        for i, piece in enumerate(pieces, start=1):
            chunks.append({
                "source_name": sec.get("source_name", "document"),
                "source_type": sec.get("source_type", "text"),
                "location": sec.get("location", "document"),
                "chunk_label": f"chunk {i}",
                "text": piece,
            })
    return chunks


__all__ = [
    "SUPPORTED_DOC_RAG_EXTENSIONS",
    "normalize_doc_text",
    "extract_pdf_sections",
    "extract_docx_sections",
    "extract_xlsx_sections",
    "extract_sections_from_uploaded_file",
    "build_chunks_from_sections",
]
